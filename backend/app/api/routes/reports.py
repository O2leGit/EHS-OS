import json
from datetime import datetime, timezone
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from app.core.database import get_db
from app.core.auth import get_current_user

router = APIRouter()


# ---------------------------------------------------------------------------
# Gap Analysis endpoints (existing)
# ---------------------------------------------------------------------------

@router.get("/gap-analysis")
async def gap_analysis(db=Depends(get_db), current_user: dict = Depends(get_current_user)):
    tid = current_user["tenant_id"]

    rows = await db.fetch(
        """SELECT framework_tier, framework_category, framework_series,
                  coverage_status, gaps, risk_severity, ai_reasoning
           FROM document_analyses WHERE tenant_id = $1::uuid
           ORDER BY framework_tier, framework_category, framework_series""",
        tid,
    )

    # Build tier/category hierarchy
    tiers_map = {}
    all_gaps = []
    covered_count = 0
    total_count = 0

    for r in rows:
        tier = r["framework_tier"] or "Unknown"
        category = r["framework_category"] or "Unknown"
        series = r["framework_series"] or ""
        status = r["coverage_status"] or "gap"

        total_count += 1
        if status == "covered":
            covered_count += 1

        if tier not in tiers_map:
            tiers_map[tier] = {}
        if category not in tiers_map[tier]:
            tiers_map[tier][category] = {
                "category": category,
                "series": series,
                "coverage_status": status,
                "gaps": [],
                "risk_severity": r["risk_severity"],
                "ai_reasoning": r["ai_reasoning"],
            }

        gap_list = r["gaps"] if isinstance(r["gaps"], list) else json.loads(r["gaps"]) if r["gaps"] else []
        if status != "covered" and gap_list:
            for gap_text in gap_list:
                gap_entry = {
                    "category": category,
                    "gap_description": gap_text,
                    "risk_severity": r["risk_severity"],
                    "recommendation": r["ai_reasoning"],
                }
                tiers_map[tier][category]["gaps"].append(gap_text)
            all_gaps.append(gap_entry)

    overall_score = round((covered_count / total_count) * 100) if total_count > 0 else 0

    tiers = []
    for tier_name, categories in tiers_map.items():
        tiers.append({
            "tier": tier_name,
            "categories": list(categories.values()),
        })

    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    top_gaps = sorted(
        all_gaps,
        key=lambda g: severity_order.get((g["risk_severity"] or "low").lower(), 4),
    )[:10]

    return {
        "overall_score": overall_score,
        "tiers": tiers,
        "top_gaps": top_gaps,
    }


@router.get("/gap-analysis/{category}")
async def gap_analysis_detail(category: str, db=Depends(get_db), current_user: dict = Depends(get_current_user)):
    tid = current_user["tenant_id"]

    rows = await db.fetch(
        """SELECT id, framework_tier, framework_category, framework_series,
                  coverage_status, gaps, ai_reasoning, chunk_text, risk_severity,
                  created_at
           FROM document_analyses
           WHERE tenant_id = $1::uuid AND framework_category = $2
           ORDER BY created_at DESC""",
        tid, category,
    )

    return [dict(r) for r in rows]


# ---------------------------------------------------------------------------
# AI Executive Reports
# ---------------------------------------------------------------------------

class ReportRequest(BaseModel):
    report_type: str  # weekly, monthly, quarterly, annual


async def gather_report_data(db, tid):
    """Gather all tenant data for report generation."""
    # Incidents
    incidents = await db.fetch(
        """SELECT incident_number, incident_type, severity, title, description,
                  location, status, created_at
           FROM incidents WHERE tenant_id=$1 ORDER BY created_at DESC""", tid)

    # CAPAs
    capas = await db.fetch(
        """SELECT c.capa_number, c.title, c.capa_type, c.status, c.priority,
                  c.due_date, c.created_at, u.full_name as assigned_to
           FROM capas c LEFT JOIN users u ON c.assigned_to = u.id
           WHERE c.tenant_id=$1 ORDER BY c.created_at DESC""", tid)

    # Sites
    sites = await db.fetch(
        "SELECT name, code, site_type, employee_count FROM sites WHERE tenant_id=$1 AND is_active=true", tid)

    # Coverage
    coverage = await db.fetch(
        "SELECT framework_tier, framework_category, coverage_status, gaps FROM document_analyses WHERE tenant_id=$1", tid)

    # Audit readiness factors
    cov_rows = await db.fetch(
        "SELECT coverage_status, COUNT(*) as cnt FROM document_analyses WHERE tenant_id=$1 GROUP BY coverage_status", tid)
    cov_total = sum(r["cnt"] for r in cov_rows)
    cov_covered = sum(r["cnt"] for r in cov_rows if r["coverage_status"] == "covered")
    cov_partial = sum(r["cnt"] for r in cov_rows if r["coverage_status"] == "partial")
    coverage_score = int(((cov_covered + cov_partial * 0.5) / max(cov_total, 1)) * 100)

    overdue_capas = await db.fetchval(
        "SELECT COUNT(*) FROM capas WHERE tenant_id=$1 AND status != 'closed' AND due_date < CURRENT_DATE", tid)
    open_capas = await db.fetchval(
        "SELECT COUNT(*) FROM capas WHERE tenant_id=$1 AND status IN ('open','in_progress')", tid)
    total_capas = await db.fetchval("SELECT COUNT(*) FROM capas WHERE tenant_id=$1", tid)

    total_inc = await db.fetchval("SELECT COUNT(*) FROM incidents WHERE tenant_id=$1", tid)
    closed_inc = await db.fetchval(
        "SELECT COUNT(*) FROM incidents WHERE tenant_id=$1 AND status='closed'", tid)
    near_misses = await db.fetchval(
        "SELECT COUNT(*) FROM incidents WHERE tenant_id=$1 AND incident_type='near_miss'", tid)
    injuries = await db.fetchval(
        "SELECT COUNT(*) FROM incidents WHERE tenant_id=$1 AND incident_type='injury'", tid)

    total_employees = sum(s["employee_count"] or 0 for s in sites)
    hours_worked = total_employees * 2000 * 0.5
    total_recordable = await db.fetchval(
        "SELECT COUNT(*) FROM incidents WHERE tenant_id=$1 AND incident_type IN ('injury','illness')", tid)
    trir = round((total_recordable / max(hours_worked, 1)) * 200000, 2) if hours_worked > 0 else 0

    def serialize(rows):
        result = []
        for r in rows:
            d = dict(r)
            for k, v in d.items():
                if hasattr(v, 'isoformat'):
                    d[k] = v.isoformat()
            result.append(d)
        return result

    return {
        "incidents": serialize(incidents),
        "capas": serialize(capas),
        "sites": [dict(s) for s in sites],
        "coverage": serialize(coverage),
        "metrics": {
            "total_incidents": total_inc,
            "closed_incidents": closed_inc,
            "investigation_closure_rate": round(closed_inc / max(total_inc, 1) * 100, 1),
            "near_misses": near_misses,
            "injuries": injuries,
            "near_miss_ratio": f"{near_misses}:{max(injuries, 1)}",
            "open_capas": open_capas,
            "overdue_capas": overdue_capas,
            "total_capas": total_capas,
            "framework_coverage_pct": coverage_score,
            "total_employees": total_employees,
            "trir": trir,
            "total_sites": len(sites),
        }
    }


def build_report_prompt(report_type: str, data: dict) -> str:
    period_descriptions = {
        "weekly": "the past week",
        "monthly": "the past month",
        "quarterly": "the past quarter (3 months)",
        "annual": "the past year"
    }
    period = period_descriptions.get(report_type, "the reporting period")

    return f"""You are a senior EHS Director writing a formal executive report for leadership. Write in first person as a seasoned EHS professional. Be direct, data-driven, and action-oriented. No disclaimers, no hedging.

Generate a comprehensive {report_type} EHS Executive Report covering {period}.

Use this data from the EHS management system:

INCIDENTS: {json.dumps(data['incidents'][:30], default=str)}

CAPAs: {json.dumps(data['capas'], default=str)}

SITES: {json.dumps(data['sites'], default=str)}

FRAMEWORK COVERAGE: {json.dumps(data['coverage'][:20], default=str)}

KEY METRICS: {json.dumps(data['metrics'], default=str)}

Write the report in this JSON structure:
{{
    "executive_summary": "2-3 paragraph executive summary with key findings, written as an EHS director would present to the C-suite",
    "sections": [
        {{
            "title": "Section Title",
            "content": "Rich paragraph content with specific data references, incident numbers, and analysis",
            "metrics": [{{"label": "Metric Name", "value": "value", "trend": "up|down|stable", "status": "good|warning|critical"}}],
            "recommendations": ["Specific actionable recommendation with timeline"]
        }}
    ],
    "risk_assessment": {{
        "overall_level": "low|moderate|elevated|high|critical",
        "summary": "Overall risk posture assessment",
        "top_risks": [
            {{"risk": "Description", "severity": "critical|high|medium|low", "mitigation": "What to do"}}
        ]
    }},
    "action_items": [
        {{"action": "Specific action", "owner": "Role/Name", "deadline": "Timeframe", "priority": "critical|high|medium|low"}}
    ],
    "looking_ahead": "Forward-looking paragraph on priorities for next period"
}}

Required sections based on report type:
- Weekly: Safety Performance Summary, Incident Review, CAPA Status, Priorities for Next Week
- Monthly: Safety Metrics & Trends, Incident Analysis, CAPA Performance, Compliance Status, Recommendations
- Quarterly: Executive Dashboard, Trend Analysis (TRIR/DART), Site Benchmarking, Framework Compliance, Strategic Recommendations, Budget Impact
- Annual: Year in Review, Annual Safety Statistics, Compliance Achievement, Program Maturity Assessment, Strategic Plan for Next Year

Include specific incident numbers (INC-XXXX), CAPA references (CAPA-XXXX), and site names. Use the actual data provided.
Respond ONLY with valid JSON. No markdown, no code blocks."""


@router.post("/generate")
async def generate_report(req: ReportRequest, db=Depends(get_db), user: dict = Depends(get_current_user)):
    tid = user["tenant_id"]

    if req.report_type not in ("weekly", "monthly", "quarterly", "annual"):
        raise HTTPException(400, "report_type must be weekly, monthly, quarterly, or annual")

    # Gather data
    data = await gather_report_data(db, tid)

    # Generate via Claude
    import traceback
    try:
        from app.services.ai import get_client
        client = get_client()

        prompt = build_report_prompt(req.report_type, data)

        response = await client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )

        content_text = response.content[0].text
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, f"AI generation failed: {str(e)}")
    try:
        if "```json" in content_text:
            content_text = content_text.split("```json")[1].split("```")[0]
        elif "```" in content_text:
            content_text = content_text.split("```")[1].split("```")[0]
        report_content = json.loads(content_text.strip())
    except (json.JSONDecodeError, IndexError):
        report_content = {"executive_summary": content_text, "sections": [], "risk_assessment": {}, "action_items": [], "looking_ahead": ""}

    # Build title
    now = datetime.now(timezone.utc)
    type_labels = {"weekly": "Weekly", "monthly": "Monthly", "quarterly": "Quarterly", "annual": "Annual"}
    title = f"{type_labels[req.report_type]} EHS Executive Report - {now.strftime('%B %d, %Y')}"
    period_label = now.strftime("%B %d, %Y")

    # Save to database
    report = await db.fetchrow(
        """INSERT INTO reports (tenant_id, report_type, period_label, title, content, created_by)
           VALUES ($1::uuid, $2, $3, $4, $5::jsonb, $6::uuid) RETURNING id, created_at""",
        tid, req.report_type, period_label, title, json.dumps(report_content), user["sub"]
    )

    return {
        "id": str(report["id"]),
        "title": title,
        "report_type": req.report_type,
        "period_label": period_label,
        "content": report_content,
        "created_at": report["created_at"].isoformat(),
    }


@router.get("/")
async def list_reports(db=Depends(get_db), user: dict = Depends(get_current_user)):
    tid = user["tenant_id"]
    rows = await db.fetch(
        """SELECT r.id, r.report_type, r.period_label, r.title, r.created_at, u.full_name as created_by
           FROM reports r LEFT JOIN users u ON r.created_by = u.id
           WHERE r.tenant_id=$1::uuid ORDER BY r.created_at DESC""", tid)
    return [
        {
            "id": str(r["id"]),
            "report_type": r["report_type"],
            "period_label": r["period_label"],
            "title": r["title"],
            "created_by": r["created_by"],
            "created_at": r["created_at"].isoformat(),
        }
        for r in rows
    ]


@router.get("/{report_id}/download")
async def download_report(report_id: str, db=Depends(get_db), user: dict = Depends(get_current_user)):
    """Download report as Word document."""
    tid = user["tenant_id"]
    row = await db.fetchrow(
        """SELECT r.id, r.report_type, r.period_label, r.title, r.content, r.created_at, u.full_name as created_by
           FROM reports r LEFT JOIN users u ON r.created_by = u.id
           WHERE r.id=$1::uuid AND r.tenant_id=$2::uuid""", report_id, tid)
    if not row:
        raise HTTPException(404, "Report not found")

    content = row["content"]
    if isinstance(content, str):
        content = json.loads(content)

    # Generate Word document
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(row["title"], level=0)

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Report Type: ").bold = True
    meta.add_run(f"{row['report_type'].title()}\n")
    meta.add_run(f"Period: ").bold = True
    meta.add_run(f"{row['period_label']}\n")
    meta.add_run(f"Generated: ").bold = True
    meta.add_run(f"{row['created_at'].strftime('%B %d, %Y at %I:%M %p')}\n")
    if row["created_by"]:
        meta.add_run(f"Prepared by: ").bold = True
        meta.add_run(f"{row['created_by']}")

    doc.add_paragraph()  # spacer

    # Executive Summary
    if content.get("executive_summary"):
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(content["executive_summary"])

    # Sections
    for section in content.get("sections", []):
        doc.add_heading(section.get("title", ""), level=1)
        if section.get("content"):
            doc.add_paragraph(section["content"])

        # Metrics table
        metrics = section.get("metrics", [])
        if metrics:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"
            hdr[2].text = "Trend"
            hdr[3].text = "Status"
            for m in metrics:
                row_cells = table.add_row().cells
                row_cells[0].text = str(m.get("label", ""))
                row_cells[1].text = str(m.get("value", ""))
                row_cells[2].text = str(m.get("trend", ""))
                row_cells[3].text = str(m.get("status", ""))
            doc.add_paragraph()  # spacer

        # Recommendations
        recs = section.get("recommendations", [])
        if recs:
            doc.add_heading("Recommendations", level=2)
            for rec in recs:
                doc.add_paragraph(rec, style='List Bullet')

    # Risk Assessment
    risk = content.get("risk_assessment", {})
    if risk:
        doc.add_heading("Risk Assessment", level=1)
        if risk.get("overall_level"):
            p = doc.add_paragraph()
            p.add_run(f"Overall Risk Level: ").bold = True
            p.add_run(risk["overall_level"].upper())
        if risk.get("summary"):
            doc.add_paragraph(risk["summary"])
        for tr in risk.get("top_risks", []):
            p = doc.add_paragraph()
            p.add_run(f"[{tr.get('severity', 'medium').upper()}] ").bold = True
            p.add_run(f"{tr.get('risk', '')}")
            if tr.get("mitigation"):
                doc.add_paragraph(f"Mitigation: {tr['mitigation']}", style='List Bullet')

    # Action Items
    actions = content.get("action_items", [])
    if actions:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = "Action"
        hdr[1].text = "Owner"
        hdr[2].text = "Deadline"
        hdr[3].text = "Priority"
        for a in actions:
            row_cells = table.add_row().cells
            row_cells[0].text = str(a.get("action", ""))
            row_cells[1].text = str(a.get("owner", ""))
            row_cells[2].text = str(a.get("deadline", ""))
            row_cells[3].text = str(a.get("priority", ""))

    # Looking Ahead
    if content.get("looking_ahead"):
        doc.add_heading("Looking Ahead", level=1)
        doc.add_paragraph(content["looking_ahead"])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by EHS-OS | Managed by Parzy Consulting | Powered by ScaleOS & IkigaiOS")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Write to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{row['report_type']}_ehs_report_{row['period_label'].replace(' ', '_').replace(',', '')}.docx"

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/{report_id}")
async def get_report(report_id: str, db=Depends(get_db), user: dict = Depends(get_current_user)):
    tid = user["tenant_id"]
    row = await db.fetchrow(
        """SELECT r.id, r.report_type, r.period_label, r.title, r.content, r.created_at, u.full_name as created_by
           FROM reports r LEFT JOIN users u ON r.created_by = u.id
           WHERE r.id=$1::uuid AND r.tenant_id=$2::uuid""", report_id, tid)
    if not row:
        raise HTTPException(404, "Report not found")

    content = row["content"]
    if isinstance(content, str):
        content = json.loads(content)

    return {
        "id": str(row["id"]),
        "report_type": row["report_type"],
        "period_label": row["period_label"],
        "title": row["title"],
        "content": content,
        "created_by": row["created_by"],
        "created_at": row["created_at"].isoformat(),
    }
